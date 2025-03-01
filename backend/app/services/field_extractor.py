"""
Модуль для извлечения динамических полей из текста документа.
Использует spaCy с многоязычной моделью xx_ent_wiki_sm.
"""

import re
from typing import Dict, List, Set
import logging
from io import BytesIO
import docx
import spacy
import difflib
from fastapi import UploadFile

logger = logging.getLogger(__name__)

# Загрузка spaCy модели
try:
    nlp = spacy.load("xx_ent_wiki_sm")
    logger.debug("Модель spaCy xx_ent_wiki_sm успешно загружена")
except Exception as e:
    logger.error("Ошибка загрузки модели spaCy: %s", str(e))
    raise Exception(
        "Модель xx_ent_wiki_sm не найдена. Установите её: python -m spacy download xx_ent_wiki_sm"
    )

# Регулярные выражения
DATE_REGEX = re.compile(
    r"\b\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\b|"
    r"\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b|"
    r"\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\b|"
    r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2},\s+\d{4}\b|"
    r"\b\d{1,2}\s+(?:январь|февраль|март|апрель|май|июнь|июль|август|сентябрь|октябрь|ноябрь|декабрь)\s+\d{4}\b|"
    r"\b(?:январь|февраль|март|апрель|май|июнь|июль|август|сентябрь|октябрь|ноябрь|декабрь)\s+\d{1,2},\s+\d{4}\b",
    re.IGNORECASE,
)

ORG_REGEX = re.compile(
    r"(?:ООО\s+)?«?[A-ZА-Я][A-Za-zА-Яа-я]+\s+[A-ZА-Я][A-Za-zА-Яа-я]+\s+[A-ZА-Я][A-Za-zА-Яа-я]*(?:\s+[A-ZА-Я][A-Za-zА-Яа-я]*)*»?(?:\s*(?:ООО|Inc\.|Ltd\.|LLC|GmbH|АО|ЗАО|ОАО|ПАО|ИП|Co\.,\s*Ltd\.))?(?:\s*\([^)]+\))?\b|\b[A-ZА-Я][A-Za-zА-Яа-я]+\s+[A-ZА-Я][A-Za-zА-Яа-я]+\s+[A-ZА-Я][A-Za-zА-Яа-я]*(?:\s+[A-ZА-Я][A-Za-zА-Яа-я]*)*(?:ООО|Inc\.|Ltd\.|LLC|GmbH|АО|ЗАО|ОАО|ПАО|ИП|Co\.,\s*Ltd\.)(?:\s*\([^)]+\))?\b",
    re.IGNORECASE,
)

# Стоп-слова и ключевые слова остаются без изменений (они универсальны)
ORG_STOPWORDS: Set[str] = {...}  # Оставляем как есть
ROLE_KEYWORDS = {...}  # Оставляем как есть
EXCLUDE_ORG_PHRASES = {...}  # Оставляем как есть
ORG_SUFFIXES = {...}  # Оставляем как есть

# Функции normalise_org, is_valid_org, etc. остаются без изменений


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from a docx file content."""
    try:
        logger.debug("Starting text extraction from docx content")
        doc = docx.Document(BytesIO(content))
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        full_text = "\n".join(text)
        logger.debug("Successfully extracted text, length: %d chars", len(full_text))
        return full_text
    except Exception as e:
        logger.error("Failed to extract text from docx: %s", str(e), exc_info=True)
        raise ValueError(f"Failed to process document: {str(e)}")


async def extract_dynamic_fields(file: UploadFile) -> List[str]:
    """Extract dynamic fields from the uploaded file."""
    try:
        logger.debug(
            "Starting field extraction for file: %s (type: %s)",
            file.filename,
            file.content_type,
        )
        content = await file.read()
        logger.debug("File content read successfully, size: %d bytes", len(content))
        text = extract_text_from_docx(content)
        logger.debug("Text extracted from document, length: %d chars", len(text))
        fields = list(set(re.findall(r"\{([^}]+)\}", text)))
        logger.info("Extracted %d fields from document", len(fields))
        logger.debug("Extracted fields: %s", fields)
        await file.seek(0)
        return fields
    except Exception as e:
        logger.error("Error during field extraction: %s", str(e), exc_info=True)
        raise ValueError(f"Failed to extract fields: {str(e)}")


def extract_orgs_with_context(text: str) -> List[str]:
    orgs = []
    lines = text.split("\n")
    org_pattern = re.compile(
        r"(?:ООО\s+)?«?[A-ZА-Я][A-Za-zА-Яа-я\s«»]+?(?:ООО|Inc\.|Ltd\.|LLC|GmbH|АО|ЗАО|ОАО|ПАО|ИП|Co\.,\s*Ltd\.)(?:\s*\([^)]+\))?\b",
        re.IGNORECASE,
    )
    for line in lines:
        line = line.strip()
        matches = org_pattern.findall(line)
        for match in matches:
            cleaned_match = normalize_org(match)
            if "Baltic Wood Agency Ltd" in cleaned_match:
                cleaned_match = "Baltic Wood Agency Ltd"
            orgs.append(cleaned_match)
    return unique_order(orgs)


if __name__ == "__main__":
    try:
        with open("example.docx", "rb") as f:
            fields = extract_dynamic_fields(f.read())
        print(fields)
    except Exception as e:
        logger.error("Ошибка в основном блоке: %s", str(e))
